import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def is_arabic(text):
    """Checks if a string contains any Arabic characters."""
    if not isinstance(text, str):
        return False
    for char in text:
        if '\u0600' <= char <= '\u06FF':
            return True
    return False

def set_rtl(paragraph):
    """Sets paragraph direction to Right-to-Left."""
    pPr = paragraph._p.get_or_add_pPr()
    
    bidi_element = pPr.find(qn('w:bidi'))
    if bidi_element is None:
        bidi_element = OxmlElement('w:bidi')
        pPr.append(bidi_element)
        
    bidi_element.set(qn('w:val'), '1')

def set_cell_direction(cell, direction="rtl"):
    """Sets cell content direction."""
    tcPr = cell._tc.get_or_add_tcPr()
    textDirection = OxmlElement('w:textDirection')
    if direction == "rtl":
        textDirection.set(qn('w:val'), "btLr")
    else:
        textDirection.set(qn('w:val'), "lrTb")
    tcPr.append(textDirection)

def generate_bilingual_profile_doc(data, labels_en, labels_ar, attachments_meta_en, attachments_meta_ar, specific_labels_en, specific_labels_ar, doc_type='member'):
    """Generates a bilingual Word document for a member or club profile."""
    try:
        document = Document()
        style = document.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        style._element.rPr.rFonts.set(qn('w:cs'), 'Arial')

        title_en = "Member Profile" if doc_type == 'member' else "Club Profile"
        title_ar = "ملف عضو" if doc_type == 'member' else "ملف هيئة"
        p = document.add_paragraph()
        p.add_run(f'{title_en} - {title_ar}').bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        document.add_paragraph()

        if doc_type == 'member' and data.get('photo_path') and os.path.exists(data['photo_path']):
            try:
                document.add_picture(data['photo_path'], width=Inches(1.5))
                document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Could not add picture to Word doc: {e}")

        # --- Main Information Table ---
        document.add_heading('Basic Information - المعلومات الأساسية', level=1)
        table = document.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        widths = (Inches(2.1), Inches(2.3), Inches(2.1))
        for i, width in enumerate(widths):
            table.columns[i].width = width

        # Header Row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Details - التفاصيل'
        hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        hdr_p_ar_field = hdr_cells[2].paragraphs[0]
        hdr_p_ar_field.text = 'الحقل'
        hdr_p_ar_field.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_rtl(hdr_p_ar_field)

        for key, label_en in labels_en.items():
            label_ar = labels_ar.get(key, label_en)
            value = data.get(key, 'N/A')
            if not value: value = 'N/A'

            row_cells = table.add_row().cells
            # Column 1: English Field
            row_cells[0].text = label_en
            
            # Column 2: Details (Single language based on input)
            p_details = row_cells[1].paragraphs[0]
            p_details.text = str(value)
            if is_arabic(value):
                set_rtl(p_details)
                p_details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Column 3: Arabic Field
            p_ar_label = row_cells[2].paragraphs[0]
            p_ar_label.text = label_ar
            p_ar_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(p_ar_label)

        # --- Specific Data ---
        specific_data_key = 'specific_data' if doc_type == 'member' else 'attachments_data'
        specific_data = data.get(specific_data_key, '{}')
        if isinstance(specific_data, str):
            specific_data = json.loads(specific_data)

        non_attachment_keys = {k: v for k, v in specific_data.items() if not k.endswith(('_docs', '_certs', '_receipts', '_license'))}
        if non_attachment_keys:
            document.add_heading('Specialization Details - تفاصيل التخصص', level=1)
            spec_table = document.add_table(rows=1, cols=3)
            spec_table.style = 'Table Grid'
            spec_table.autofit = False
            for i, width in enumerate(widths):
                spec_table.columns[i].width = width
            
            # Header Row
            spec_hdr_cells = spec_table.rows[0].cells
            spec_hdr_cells[0].text = 'Field'
            spec_hdr_cells[1].text = 'Details - التفاصيل'
            spec_hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

            spec_hdr_p_ar_field = spec_hdr_cells[2].paragraphs[0]
            spec_hdr_p_ar_field.text = 'الحقل'
            spec_hdr_p_ar_field.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_rtl(spec_hdr_p_ar_field)

            for key, value in non_attachment_keys.items():
                label_en = specific_labels_en.get(key, key.replace('_', ' ').title())
                label_ar = specific_labels_ar.get(key, key.replace('_', ' ').title())
                
                row_cells = spec_table.add_row().cells
                # Column 1: English Field
                row_cells[0].text = label_en

                # Column 2: Details (Single language based on input)
                p_details = row_cells[1].paragraphs[0]
                p_details.text = str(value)
                if is_arabic(str(value)):
                    set_rtl(p_details)
                    p_details.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Column 3: Arabic Field
                p_ar_label = row_cells[2].paragraphs[0]
                p_ar_label.text = label_ar
                p_ar_label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                set_rtl(p_ar_label)

        # --- Attachments ---
        attachment_keys = [k for k in specific_data if k.endswith(('_docs', '_certs', '_receipts', '_license'))]
        if attachment_keys:
            document.add_heading('Attachments - المرفقات', level=1)
            for key in attachment_keys:
                file_paths = specific_data.get(key)
                if file_paths:
                    label_en = attachments_meta_en.get(key, key.replace('_', ' ').title())
                    label_ar = attachments_meta_ar.get(key, label_en)
                    document.add_paragraph(f"{label_en} - {label_ar}", style='Intense Quote')
                    for path in file_paths:
                        document.add_paragraph(os.path.basename(path), style='List Bullet')

        # --- Save temporary file ---
        temp_dir = os.path.join('assets', 'temp_reports')
        os.makedirs(temp_dir, exist_ok=True)
        
        id_key = 'pkf_id' if doc_type == 'member' else 'club_membership_id'
        safe_id = "".join(c for c in data.get(id_key, 'profile') if c.isalnum()).rstrip()
        
        filename = f"{safe_id}_{doc_type}_profile.docx"
        temp_path = os.path.join(temp_dir, filename)
        document.save(temp_path)
        return temp_path

    except Exception as e:
        print(f"Error generating Word document: {e}")
        return None